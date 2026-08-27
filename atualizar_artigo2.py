# -*- coding: utf-8 -*-
"""Gera o Artigo 2 revisado e a carta de resposta a partir de metricas.json."""
from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor
from docx.oxml import OxmlElement

ROOT = Path(r"f:\git\gastrack2")
RES = ROOT / "resultados_anomalias"
M = json.loads((RES / "metricas.json").read_text(encoding="utf-8"))
FIG1 = ROOT / "_artigo2_media" / "word" / "media" / "image2.png"
OUT_ART = ROOT / "Artigo_2_GasTrack_Anomalias_Revista_Topicos_REVISADO.docx"
OUT_CARTA = ROOT / "Carta_Resposta_Revisores_Artigo2.docx"
DESKTOP = Path(r"C:\Users\joaoclaudio\OneDrive\Desktop")

MET = ["Regra robusta", "Isolation Forest", "Local Outlier Factor", "One-Class SVM"]
TIPOS = ["queda", "deriva", "ruido", "travado", "picos"]
TIPOS_L = ["Queda abrupta", "Deriva", "Ruído excessivo", "Sensor travado", "Picos"]


def br(x, n=4):
    return f"{float(x):.{n}f}".replace(".", ",")


def md(d, n=4):
    return f"{br(d['media'], n)} ± {br(d['desvio'], n)}"


def ic(d, n=4):
    return f"{br(d['media'], n)} (IC95% {br(d['ic95_lo'], n)}–{br(d['ic95_hi'], n)})"


R = M["resumo"]
PT = M["por_tipo"]
DIF = M["diff_f1_lof_regra"]
PPV = M["ppv"]
ABL = M["ablacao"]
OOD = M["ood_fpr"]
JAN = M["janela"]
SEV = M["severidade"]
CUS = M["custo"]
OVL = M["overlap"]
CNF = M["confusao_semente0"]
MCN = M["mcnemar_semente0"]
SK = M["sensibilidade_k"]
SIF = M["sensibilidade_if"]


def set_font(run, size=12, bold=False, italic=False, name="Times New Roman"):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def fmt_p(p, align="justify", before=0, after=0, first=False, line=1.5):
    if align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.first_line_indent = Cm(1.25) if first else Cm(0)


def add(doc, text, *, size=12, bold=False, italic=False, align="justify",
        before=0, after=0, first=False, line=1.5):
    p = doc.add_paragraph()
    fmt_p(p, align, before, after, first, line)
    r = p.add_run(text)
    set_font(r, size, bold, italic)
    return p


def caption(doc, text):
    add(doc, text, size=10, italic=False, align="left", before=6, after=10, first=False, line=1.15)


def picture(doc, path, width_cm=16.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Cm(width_cm))


def shade(cell, hex_color="1F4E79"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def tabela(doc, header, rows, col_cm=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Table Grid"
    for j, h in enumerate(header):
        cell = t.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, 10, True)
        r.font.color.rgb = RGBColor(255, 255, 255)
        shade(cell, "1F4E79")
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_font(r, 10, False)
            if i % 2:
                shade(cell, "E8EEF4")
    if col_cm:
        for row in t.rows:
            for j, w in enumerate(col_cm):
                row.cells[j].width = Cm(w)
    return t


def setup(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(3.0)
    sec.bottom_margin = Cm(2.0)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def construir_artigo():
    doc = Document()
    setup(doc)

    add(doc,
        "DETECÇÃO NÃO SUPERVISIONADA DE ANOMALIAS EM SÉRIES TEMPORAIS DE PRESSÃO: "
        "UM BENCHMARK SINTÉTICO DE PROVA DE CONCEITO NO CONTEXTO GASTRACK",
        size=14, bold=True, align="center", after=4, line=1.15)
    add(doc,
        "UNSUPERVISED ANOMALY DETECTION IN PRESSURE TIME SERIES: A SYNTHETIC "
        "PROOF-OF-CONCEPT BENCHMARK IN THE GASTRACK CONTEXT",
        size=13, bold=True, align="center", after=8, line=1.15)
    add(doc, "João Cláudio Nunes Carvalho", size=12, align="center", after=6, line=1.15)

    add(doc, "RESUMO", size=12, bold=True, align="left", before=8, after=2, line=1.5)
    add(doc,
        "Este trabalho apresenta um benchmark sintético de prova de conceito para "
        "detecção não supervisionada de anomalias em séries temporais de pressão no "
        "contexto do sistema GasTrack. O estudo não visa diagnosticar vazamentos: uma "
        "anomalia de pressão é um indício de comportamento incomum, compatível com "
        "vazamento, variação de consumo, falha de válvula ou defeito do próprio sensor. "
        f"Foram comparados uma regra estatística robusta, Isolation Forest, Local Outlier "
        f"Factor (LOF) e One-Class SVM sobre janelas de 24 leituras espaçadas de 5 min "
        f"(extensão temporal de 115 min, aproximadamente duas horas). O protocolo separa "
        f"5.000 janelas normais de treino, 2.000 de calibração independente do limiar "
        f"(percentil 99 dos escores) e um teste com 2.500 janelas normais e 2.500 anômalas. "
        f"O vetor possui onze atributos de tendência, dispersão, abruptividade e planicidade; "
        f"não há variável duplicada. Em 20 sementes, o LOF obteve F1 de {md(R['Local Outlier Factor']['f1'])}, "
        f"próximo da regra robusta ({md(R['Regra robusta']['f1'])}). A diferença absoluta de F1 "
        f"foi {br(DIF['media'])} (IC95% {br(DIF['ic95_lo'])}–{br(DIF['ic95_hi'])}). A vantagem do LOF "
        f"no sensor travado depende da representação, em especial do atributo de planicidade. "
        f"Com prevalência de 1%, a precisão positiva esperada do LOF cai para {br(PPV['0.01']['Local Outlier Factor'], 3)}, "
        f"o que impede interpretar o F1 do teste balanceado como desempenho de produção. "
        f"Python {M['hiperparametros']['python']} e scikit-learn {M['hiperparametros']['sklearn']} "
        f"foram utilizados; o código do gerador e da avaliação é disponibilizado como material suplementar.",
        size=12, align="justify", first=False, line=1.5)
    add(doc,
        "Palavras-chave: Detecção de anomalias; Séries temporais; Internet das Coisas; "
        "Sensores de pressão; Aprendizado de máquina.",
        size=12, italic=True, align="justify", before=6, after=8, first=False, line=1.5)

    add(doc, "ABSTRACT", size=12, bold=True, align="left", before=8, after=2, line=1.5)
    add(doc,
        "This paper presents a synthetic proof-of-concept benchmark for unsupervised "
        "anomaly detection in pressure time series in the GasTrack context. The study "
        "does not aim to diagnose leaks: a pressure anomaly is evidence of unusual "
        "behavior, which may arise from leakage, consumption changes, valve faults or "
        "sensor defects. A robust statistical rule, Isolation Forest, Local Outlier Factor "
        "(LOF) and One-Class SVM were compared on windows of 24 readings spaced by 5 min "
        "(temporal span of 115 min). The protocol separates 5,000 normal training windows, "
        "2,000 independent calibration windows for the 99th-percentile threshold, and a test "
        "set with 2,500 normal and 2,500 anomalous windows. Eleven features describe trend, "
        "dispersion, abruptness and flatness; no duplicated variable is used. Over 20 seeds, "
        f"LOF achieved an F1 of {md(R['Local Outlier Factor']['f1'])}, close to the robust rule "
        f"({md(R['Regra robusta']['f1'])}). The F1 gap was {br(DIF['media'])} "
        f"(95% CI {br(DIF['ic95_lo'])}–{br(DIF['ic95_hi'])}). LOF’s advantage on stuck-sensor "
        "windows depends on the representation, especially the flatness feature. At 1% "
        f"prevalence, the expected positive predictive value of LOF falls to "
        f"{br(PPV['0.01']['Local Outlier Factor'], 3)}, so balanced-test F1 must not be read as "
        "production performance. Python "
        f"{M['hiperparametros']['python']} and scikit-learn {M['hiperparametros']['sklearn']} "
        "were used; generator and evaluation code is provided as supplementary material.",
        size=12, align="justify", first=False, line=1.5)
    add(doc,
        "Keywords: Anomaly detection; Time series; Internet of Things; Pressure sensors; "
        "Machine learning.",
        size=12, italic=True, align="justify", before=6, after=8, first=False, line=1.5)

    add(doc, "RESUMEN", size=12, bold=True, align="left", before=8, after=2, line=1.5)
    add(doc,
        "Este trabajo presenta un benchmark sintético de prueba de concepto para la "
        "detección no supervisada de anomalías en series temporales de presión en el "
        "contexto del sistema GasTrack. El estudio no pretende diagnosticar fugas: una "
        "anomalía de presión es un indicio de comportamiento inusual. Se compararon una "
        "regla estadística robusta, Isolation Forest, Local Outlier Factor y One-Class SVM "
        "en ventanas de 24 lecturas. El protocolo separa entrenamiento, calibración "
        "independiente del umbral y prueba. En 20 semillas, el LOF obtuvo un F1 de "
        f"{md(R['Local Outlier Factor']['f1'])}, próximo a la regla robusta "
        f"({md(R['Regra robusta']['f1'])}). Con prevalencia del 1%, el valor predictivo "
        f"positivo esperado del LOF cae a {br(PPV['0.01']['Local Outlier Factor'], 3)}. "
        "Los resultados caracterizan una prueba de concepto y no demuestran capacidad "
        "de diagnosticar fugas.",
        size=12, align="justify", first=False, line=1.5)
    add(doc,
        "Palabras clave: Detección de anomalías; Series temporales; Internet de las Cosas; "
        "Sensores de presión; Aprendizaje automático.",
        size=12, italic=True, align="justify", before=6, after=10, first=False, line=1.5)

    # ---------- 1
    add(doc, "1. INTRODUÇÃO", size=12, bold=True, align="left", before=12, after=6, line=1.5)
    add(doc,
        "Sistemas de monitoramento baseados em Internet das Coisas produzem séries temporais "
        "em volume e frequência incompatíveis com inspeção manual contínua. Em uma aplicação "
        "de gás, cada dispositivo pode registrar milhares de leituras de pressão por dia, "
        "tornando necessário distinguir automaticamente comportamento esperado de situações "
        "que merecem atenção. O desafio não é apenas detectar valores altos ou baixos, mas "
        "reconhecer mudanças no padrão temporal, como quedas abruptas, deriva, ruído anormal, "
        "travamento do sensor ou picos isolados.",
        first=True)
    add(doc,
        "A literatura de segurança de gás mostra que aprendizado de máquina tem sido aplicado "
        "tanto à identificação de vazamentos quanto ao diagnóstico de falhas de sensores. "
        "Yuan et al. (2025) analisaram dados IoT de monitoramento urbano e utilizaram engenharia "
        "de atributos temporais para classificar eventos de vazamento, obtendo recall de 92,86% "
        "com árvore de decisão. Tan et al. (2022) investigaram falhas em sensores de sistemas de "
        "monitoramento de gás e relataram elevada acurácia de diagnóstico ao combinar métodos "
        "probabilísticos. Esses estudos evidenciam que o próprio sistema de medição pode produzir "
        "anomalias que não correspondem a um vazamento físico.",
        first=True)
    add(doc,
        "Em dispositivos de borda, abordagens TinyML também têm sido exploradas para detectar "
        "gases e situações anômalas localmente. Tsoukas et al. (2023) apresentam um dispositivo "
        "de detecção baseado em TinyML, enquanto El Barkani et al. (2024) demonstram modelos "
        "compactos para classificação de vazamentos em hardware com recursos limitados, reportando "
        "memória e tempo de inferência. Embora tais abordagens sejam supervisionadas e utilizem "
        "modalidades distintas da pressão — no caso de El Barkani et al., câmera térmica MLX90640 "
        "e redes convolucionais —, elas sustentam a viabilidade de inferência embarcada, e não a "
        "transferência direta do modelo para o GasTrack.",
        first=True)
    add(doc,
        "O GasTrack é uma plataforma de telemetria de pressão já desenvolvida, composta por "
        "sensor industrial XDB305 (4–20 mA, faixa 0–250 bar), conversor HW-685, microcontrolador "
        "ESP32-ETH01 com ADC de 12 bits, publicação MQTT a cada 10 s e armazenamento em nuvem "
        "para visualização. A aplicação possui regras determinísticas de nível (incluindo limiar "
        "crítico de 20% da capacidade) e cooldown de 600 s para alerta de pressão crítica. O "
        "dashboard agrega séries em intervalos de 5 e 15 min. Entretanto, um alerta baseado "
        "exclusivamente em limiar responde a uma pergunta limitada: a pressão está abaixo de um "
        "valor definido? Uma camada de detecção de anomalias busca responder a outra questão: o "
        "comportamento recente difere significativamente do padrão normal esperado para aquele sinal?",
        first=True)
    add(doc,
        "A distinção é fundamental porque uma anomalia estatística não possui, por si só, "
        "interpretação causal. Uma queda abrupta pode ser compatível com vazamento, mas também "
        "pode refletir aumento intenso de consumo, troca de configuração ou erro de leitura. Da "
        "mesma forma, uma sequência praticamente constante pode indicar estabilidade real ou "
        "travamento do sensor. O detector desenvolvido neste artigo é, portanto, um anomaly "
        "detector for pressure telemetry, e não um leak detector. Modelos de aprendizado de "
        "máquina não devem substituir limites de segurança validados; funcionam como camada "
        "complementar de triagem.",
        first=True)
    add(doc,
        "A questão de pesquisa é: em um benchmark sintético controlado, quais métodos não "
        "supervisionados reconhecem diferentes mecanismos de alteração do sinal de pressão sem "
        "exigir um grande conjunto rotulado de falhas? O objetivo consiste em comparar uma regra "
        "estatística robusta, Isolation Forest, Local Outlier Factor e One-Class SVM, avaliando "
        "precisão, recall, F1-score, ROC-AUC, PR-AUC, taxa de falso positivo e sensibilidade por "
        "tipo de anomalia, com quantificação de incerteza entre realizações do gerador. O trabalho "
        "é explicitamente classificado como prova de conceito sintética: não substitui validação "
        "em séries reais, naturalmente desbalanceadas e sujeitas a mudanças operacionais.",
        first=True)

    # ---------- 2
    add(doc, "2. FUNDAMENTAÇÃO TEÓRICA", size=12, bold=True, align="left", before=12, after=6)
    add(doc, "2.1 Anomalias em redes de sensores e sistemas de gás", size=12, bold=True,
        align="left", before=8, after=4)
    add(doc,
        "Sensores distribuídos podem produzir dados incorretos por falhas físicas, comunicação, "
        "envelhecimento, interferência, alimentação ou condições ambientais. Souza et al. (2020), "
        "ao estudarem redes de sensores sem fio, ressaltam que a identificação automática de "
        "sensores anormais é necessária porque leituras errôneas podem comprometer as decisões "
        "posteriores. Em sistemas de gás, essa preocupação possui dimensão adicional de segurança, "
        "já que uma leitura incorreta pode tanto ocultar uma condição crítica quanto produzir alarmes falsos.",
        first=True)
    add(doc,
        "Tan et al. (2022) tratam explicitamente a distinção entre dado anormal e tipo de falha "
        "do sensor. O estudo mostra que o diagnóstico confiável depende de duas etapas: reconhecer "
        "que a observação é incompatível com o padrão esperado e, em seguida, classificar sua causa. "
        "Essa separação conceitual orienta o presente trabalho. A saída dos detectores avaliados é "
        "apenas um indicador de anormalidade, sem atribuir automaticamente a classe “vazamento”.",
        first=True)
    add(doc,
        "Yuan et al. (2025) demonstram que eventos de vazamento podem ser reconhecidos em séries "
        "IoT quando atributos de forma, volatilidade e tendência são construídos a partir das "
        "sequências. O resultado reforça a importância da engenharia de atributos. Em vez de "
        "fornecer apenas a pressão atual, uma janela temporal pode ser descrita por inclinação, "
        "variância das diferenças, amplitude, mudanças de segunda ordem e proporção de intervalos "
        "praticamente constantes.",
        first=True)

    add(doc, "2.2 Métodos não supervisionados de detecção", size=12, bold=True,
        align="left", before=8, after=4)
    add(doc,
        "Métodos não supervisionados são particularmente úteis quando exemplos rotulados de falha "
        "são raros ou difíceis de reproduzir com segurança. O treinamento pode utilizar "
        "predominantemente dados normais, modelando regiões de alta densidade ou mecanismos de "
        "isolamento. Esse paradigma é adequado para uma primeira camada de monitoramento do "
        "GasTrack, pois eventos reais de vazamento, falha ou perda de calibração tendem a ser "
        "muito menos frequentes que períodos normais.",
        first=True)
    add(doc,
        "O Isolation Forest parte da ideia de que anomalias são mais fáceis de isolar por "
        "particionamentos aleatórios do espaço de atributos. Observações raras tendem a alcançar "
        "folhas em profundidades menores, gerando escores de anomalia mais altos (Liu; Ting; Zhou, "
        "2008). O Local Outlier Factor, proposto por Breunig et al. (2000), compara a densidade "
        "local de uma observação com a densidade de seus vizinhos; a vizinhança é determinada, "
        "por padrão, por distância de Minkowski com p = 2 (euclidiana). A One-Class SVM busca "
        "construir uma fronteira que englobe a maior parte das observações normais em um espaço "
        "transformado por kernel, tratando pontos externos como novidade (Schölkopf et al., 2001). "
        "O parâmetro ν funciona como limite superior para a fração de erros de treino e limite "
        "inferior para a fração de vetores de suporte. Regras estatísticas robustas, por sua vez, "
        "oferecem maior interpretabilidade ao sinalizar quando uma característica se afasta "
        "fortemente da mediana histórica.",
        first=True)

    add(doc, "2.3 Detecção na borda e limites da interpretação", size=12, bold=True,
        align="left", before=8, after=4)
    add(doc,
        "Trabalhos com TinyML mostram que modelos compactos podem executar inferência em "
        "dispositivos com recursos limitados. Tsoukas et al. (2023) avaliaram cenários de fumaça "
        "e amônia, enquanto El Barkani et al. (2024) compararam arquiteturas convolucionais "
        "otimizadas para execução local e reportaram memória e tempo de inferência. Esses estudos "
        "ampliam a perspectiva de que uma futura versão do GasTrack possa realizar parte da "
        "triagem no próprio dispositivo. Entretanto, câmeras térmicas, sensores de concentração "
        "e sensores de pressão observam fenômenos distintos. Resultados de detecção de vazamento "
        "obtidos com uma modalidade não podem ser transferidos diretamente para outra. Uma "
        "anomalia de pressão é um indício físico-operacional, não uma prova de presença de gás "
        "no ambiente. AUC elevada, da mesma forma, não demonstra capacidade de diagnóstico causal.",
        first=True)

    # ---------- 3
    add(doc, "3. METODOLOGIA", size=12, bold=True, align="left", before=12, after=6)
    add(doc, "3.1 Arquitetura de referência e unidade de análise", size=12, bold=True,
        align="left", before=8, after=4)
    add(doc,
        "A arquitetura de referência segue o fluxo sensor–ESP32–telemetria–armazenamento–análise "
        "(Figura 1). A unidade de análise é uma janela de n = 24 leituras consecutivas, agregadas "
        "em intervalos de Δt = 5 min, convenção já disponível no dashboard do GasTrack. A diferença "
        "temporal entre a primeira e a última marca é (n − 1)Δt = 115 min; refere-se, portanto, a "
        "uma janela de aproximadamente duas horas, e não a 120 min exatos. A escolha de 24 amostras "
        "busca um compromisso entre detecção relativamente rápida (janela curta, menos contexto) e "
        "caracterização de tendência (janela longa, maior atraso). Janelas de 12 e 48 leituras foram "
        "avaliadas em análise complementar. Cada janela foi convertida em um vetor de atributos "
        "antes de ser apresentada aos algoritmos. O experimento principal utiliza janelas "
        "independentes; o efeito de janelas deslizantes sobrepostas é discutido à parte, porque "
        "em produção duas janelas sucessivas com passo de uma leitura compartilham 23 das 24 observações.",
        first=True)
    if FIG1.exists():
        picture(doc, FIG1, 16.2)
    caption(doc, "Figura 1: Fluxo lógico para detecção de anomalias em séries de pressão. Fonte: Elaborado pelo autor (2026).")

    add(doc, "3.2 Gerador sintético do comportamento normal", size=12, bold=True,
        align="left", before=8, after=4)
    add(doc,
        "As faixas de pressão e de consumo são sintéticas, inspiradas na escala do sensor XDB305 "
        "(0–250 bar) e em regimes operacionais plausíveis de cilindros industriais, e não foram "
        "ajustadas a um conjunto de campo rotulado. Declaramos explicitamente que uma janela "
        "sintética não representa de forma direta o comportamento físico de um ponto GasTrack em "
        "produção. A pressão inicial p₀ é sorteada de Uniforme(60; 145) bar, faixa intermediária "
        "da escala do sensor. A taxa base de queda r é Uniforme(0,5; 5,0) bar/h. Há modulação "
        "temporal de ±10% com período Uniforme(4; 12) h e fase uniforme em [0, 2π). Oscilação "
        "periódica tem amplitude Uniforme(0; 1,5) bar e período Uniforme(2; 8) h. O ruído de "
        "medição é gaussiano i.i.d. com desvio-padrão Uniforme(0,15; 0,80) bar. Se tᵢ = iΔt, "
        "i = 0, …, 23, a janela normal é",
        first=True)
    add(doc,
        "p(tᵢ) = p₀ − Σ_{k=0}^{i−1} r_k Δt + a sen(2π tᵢ/T + φ) + εᵢ,    εᵢ ~ N(0, σ²),",
        size=12, italic=True, align="center", before=6, after=6, first=False, line=1.15)
    add(doc,
        "com r_k = r [1 + 0,10 sen(2π t_k/T_m + φ_m)]. Todos os sorteios são independentes. O "
        "gerador de números pseudoaleatórios é numpy.random.default_rng(semente), com sementes "
        "inteiras 0 a 19 no experimento principal.",
        first=True)

    add(doc, "3.3 Mecanismos de anomalia", size=12, bold=True, align="left", before=8, after=4)
    add(doc,
        "Cada janela anômala parte de uma realização normal à qual se aplica um único mecanismo. "
        "No protocolo nominal: (i) queda abrupta — a partir de um índice j sorteado uniformemente "
        "em {12, …, n−4} (segunda metade da janela), subtrai-se Δ ~ Uniforme(7; 25) bar de todas "
        "as leituras p_j, …, p_{n−1}; (ii) deriva — adiciona-se uma queda quadrática que acumula "
        "D ~ Uniforme(8; 20) bar ao final da janela, p̃ᵢ = pᵢ − D (tᵢ/t_{n−1})²; (iii) ruído "
        "excessivo — soma-se ruído gaussiano adicional com desvio Uniforme(2,5; 6,0) bar; "
        "(iv) sensor travado — as últimas k = 8 leituras (~40 min) são substituídas por "
        "p_{n−k} + η, η ~ N(0; 0,02²), isto é, praticamente constantes, com variância residual "
        "inferior a uma contagem do ADC de 12 bits (LSB ≈ 0,061 bar); (v) picos — sorteiam-se "
        "N_s ∈ {1, 2, 3} índices distintos em {1, …, n−2}, sinais em {−1, +1} com igual "
        "probabilidade e amplitudes Uniforme(8; 25) bar. A opção por anomalias sintéticas é "
        "metodológica e de segurança: não foram criados vazamentos físicos nem utilizados rótulos "
        "inexistentes de produção. As classes representam mecanismos de alteração do sinal, não "
        "causas definitivas.",
        first=True)
    add(doc,
        "Para avaliar o limite entre normal e anormal, replicou-se o protocolo em três níveis de "
        "severidade (baixa, média e alta). Na intensidade baixa, por exemplo, a queda está em "
        "1–3 bar, a deriva acumula 2–5 bar, o ruído adicional é 1,0–1,8 bar, o travamento cobre "
        "apenas 4 leituras com residual 0,05 bar, e há um único pico de 2,5–5,0 bar. Também foi "
        "gerado um conjunto normal fora da distribuição de treino (OOD), ainda operacionalmente "
        "legítimo: pressão-base Uniforme(40; 58) bar, ou taxa Uniforme(5,5; 8,0) bar/h, ou ruído "
        "Uniforme(0,85; 1,20) bar. Esse cenário estima falsos alertas em mudanças operacionais "
        "benignas, não a generalização IID.",
        first=True)
    picture(doc, RES / "figura2_padroes.png", 16.2)
    caption(doc, "Figura 2: Exemplos de comportamento normal e dos cinco mecanismos sintéticos em uma janela de 24 leituras. Fonte: Elaborado pelo autor (2026).")

    add(doc, "3.4 Engenharia de atributos", size=12, bold=True, align="left", before=8, after=4)
    add(doc,
        "Cada janela p = (p₀, …, p₂₃) é representada por onze atributos, sem duplicação. Sejam "
        "Δpᵢ = pᵢ − pᵢ−₁ e Δ²pᵢ = Δpᵢ − Δpᵢ−₁. Unidades em bar ou bar/h, conforme o caso.",
        first=True)
    add(doc,
        "(1) inclinação OLS: β̂₁ da regressão pᵢ = β₀ + β₁ tᵢ, em bar/h; (2) inclinação de "
        "Theil–Sen: mediana das inclinações par a par i < j, β_TS = mediana{(pⱼ − pᵢ)/(tⱼ − tᵢ)}, "
        "em bar/h; (3) desvio-padrão amostral de p; (4) desvio-padrão amostral de Δp; (5) média "
        "de |Δp|; (6) máximo de |Δp|; (7) amplitude p_max − p_min; (8) mediana de |Δp|; "
        "(9) média de |Δ²p|; (10) planicidade: proporção de |Δpᵢ| < 0,08 bar (cerca de 1,3 LSB "
        "do ADC); (11) mudança líquida p₂₃ − p₀. Não se inclui cópia da amplitude. O limiar de "
        "0,08 bar é uma definição a priori alinhada à resolução do sensor, não um parâmetro "
        "otimizado no teste. O atributo de planicidade é deliberadamente informativo para o "
        "sensor travado; essa dependência é quantificada por ablação.",
        first=True)
    add(doc,
        "A padronização (média e desvio-padrão) foi ajustada exclusivamente no treino e aplicada "
        "sem novo ajuste à calibração e ao teste, apenas para LOF e One-Class SVM. Isolation "
        "Forest permaneceu na escala original, porque árvores são pouco sensíveis a transformações "
        "lineares de escala. A regra robusta utiliza mediana e desvio absoluto mediano (MAD) "
        "estimados no treino.",
        first=True)
    picture(doc, RES / "figura5_correlacao.png", 14.5)
    caption(doc, "Figura 3: Matriz de correlação de Pearson dos onze atributos no conjunto de treino (semente 0). Fonte: Elaborado pelo autor (2026).")

    add(doc, "3.5 Protocolo de treino, calibração e decisão", size=12, bold=True,
        align="left", before=8, after=4)
    add(doc,
        "Para cada semente geram-se 5.000 janelas normais de treino, 2.000 janelas normais "
        "independentes de calibração e um teste com 2.500 janelas normais e 2.500 anômalas "
        "(500 de cada mecanismo). Nenhuma anomalia entra na definição do limiar. Os quatro "
        "detectores produzem um escore contínuo cuja orientação foi unificada: valores maiores "
        "indicam maior evidência de anomalia. Na regra robusta, o escore é o máximo dos "
        "|z|_j = |x_j − mediana_j| / MAD_j. No Isolation Forest e no LOF utiliza-se "
        "−score_samples; na One-Class SVM, −decision_function. O limiar operacional é o "
        "percentil 99 desse escore no conjunto normal de calibração, visando taxa nominal de "
        "falso alarme de 1%.",
        first=True)
    add(doc,
        "O parâmetro contamination do scikit-learn foi deixado em “auto” e não participa da "
        "decisão final. Isolation Forest e LOF usam contamination internamente apenas para um "
        "offset_ de predict(); como a decisão binária é reconstruída pelo percentil 99 externo, "
        "esse parâmetro seria redundante. No LOF, novelty=True: score_samples e decision_function "
        "nunca são aplicados ao próprio conjunto de ajuste, em conformidade com a documentação "
        "do algoritmo, porque a geometria da vizinhança de um ponto já presente no treino não "
        "coincide com a de uma observação nova.",
        first=True)
    add(doc,
        "Hiperparâmetros foram definidos a priori, sem uso do conjunto de teste. Isolation Forest: "
        "300 árvores, max_samples=“auto” (min(256, n) no scikit-learn), random_state igual à "
        "semente da repetição. LOF: k = 35 vizinhos, métrica euclidiana; 35 é o ponto médio do "
        "intervalo 10–50 recomendado por Breunig et al. (2000) para MinPts, não um valor escolhido "
        "após inspecionar o teste. One-Class SVM: kernel=“rbf”, ν = 0,01 (orçamento de 1% de "
        "falso alarme, coerente com a interpretação de ν como limite superior da fração de erros "
        "de treino) e gamma=“scale”, isto é, γ = 1 / (n_atributos · Var(X)). Análise de "
        "sensibilidade posterior avaliou k ∈ {10, 20, 35, 50, 75} e n_estimators ∈ {100, 300, 500} "
        "apenas para verificar estabilidade do ranqueamento, sem retreinar a configuração principal.",
        first=True)
    add(doc,
        f"O experimento foi executado em Python {M['hiperparametros']['python']}, scikit-learn "
        f"{M['hiperparametros']['sklearn']}, NumPy {M['hiperparametros']['numpy']} e pandas "
        f"{M['hiperparametros']['pandas']}. O código do gerador, as sementes, a tabela de "
        "hiperparâmetros e o script de avaliação acompanham o manuscrito como material suplementar.",
        first=True)

    add(doc, "3.6 Métricas e incerteza", size=12, bold=True, align="left", before=8, after=4)
    add(doc,
        "Reportam-se precisão, recall, F1, ROC-AUC, PR-AUC (Average Precision), FPR, "
        "especificidade e alarmes falsos por 1.000 janelas normais. Como o teste está "
        "artificialmente balanceado (prevalência 50%), a precisão desse conjunto não é estimativa "
        "da precisão em produção. A precisão positiva esperada (PPV) em prevalência π é "
        "calculada por π·R / [π·R + FPR·(1 − π)], com R e FPR médios do teste. Média, "
        "desvio-padrão, mediana e intervalo de confiança de 95% (normal, sobre as 20 sementes) "
        "são apresentados. A diferença pareada de F1 entre LOF e regra robusta, avaliada nas "
        "mesmas janelas, é resumida por IC95%. Na semente 0 aplica-se o teste de McNemar às "
        "decisões binárias. Métricas são por janela; métricas por evento (detecção do episódio, "
        "atraso, alarmes duplicados) são ilustradas no experimento de janela deslizante e não "
        "substituem o benchmark principal.",
        first=True)

    # ---------- 4
    add(doc, "4. RESULTADOS E DISCUSSÃO", size=12, bold=True, align="left", before=12, after=6)
    add(doc, "4.1 Comparação global", size=12, bold=True, align="left", before=8, after=4)
    add(doc,
        f"A Tabela 1 resume o desempenho em 20 sementes. O LOF obteve o maior F1 médio "
        f"({md(R['Local Outlier Factor']['f1'])}), seguido da One-Class SVM "
        f"({md(R['One-Class SVM']['f1'])}) e da regra robusta ({md(R['Regra robusta']['f1'])}). "
        f"O Isolation Forest ficou atrás pelo recall ({md(R['Isolation Forest']['recall'])}). "
        f"A diferença de F1 entre LOF e regra robusta foi {ic(DIF)}, portanto sistematicamente "
        f"positiva neste gerador, mas pequena em termos absolutos ({br(100*DIF['media'], 2)} "
        f"pontos percentuais). Na semente 0, o teste de McNemar nas decisões binárias rejeitou "
        f"igualdade (n₀₁ = {MCN['n01']}, n₁₀ = {MCN['n10']}, p ≈ {MCN['p']:.1e}). A mensagem "
        f"científica mais defensável não é “o LOF venceu”, e sim que atributos temporais simples "
        f"permitem reconhecer diferentes mecanismos sintéticos, com vantagem local no padrão de "
        f"baixa dinâmica do sensor travado e com a regra robusta competitiva e mais interpretável.",
        first=True)
    add(doc,
        f"O FPR no teste normal ficou próximo do alvo de 1% "
        f"(LOF {md(R['Local Outlier Factor']['fpr'])}, equivalentes a "
        f"{md(R['Local Outlier Factor']['alarmes_por_1000'], 1)} alarmes falsos por 1.000 janelas). "
        f"A especificidade correspondente do LOF foi {md(R['Local Outlier Factor']['especificidade'])}. "
        f"ROC-AUC e PR-AUC permaneceram elevados no teste balanceado; a PR-AUC é a métrica de "
        f"ranking mais pertinente quando a classe positiva será rara em produção, mas ainda assim "
        f"não substitui o cálculo de PPV sob prevalência realista (Seção 4.4).",
        first=True)

    tabela(doc,
           ["Método", "Precisão", "Recall", "F1", "ROC-AUC", "PR-AUC", "FPR"],
           [[m,
             md(R[m]["precisao"]),
             md(R[m]["recall"]),
             md(R[m]["f1"]),
             md(R[m]["roc_auc"]),
             md(R[m]["pr_auc"]),
             md(R[m]["fpr"])] for m in MET])
    caption(doc, "Tabela 1: Métricas globais (média ± desvio-padrão em 20 sementes) no teste sintético balanceado. Fonte: Elaborado pelo autor (2026).")
    picture(doc, RES / "figura3_desempenho.png", 15.5)
    caption(doc, "Figura 4: F1-score e recall médios com desvio-padrão entre sementes. Fonte: Elaborado pelo autor (2026).")

    add(doc, "4.2 Sensibilidade por mecanismo", size=12, bold=True, align="left", before=8, after=4)
    add(doc,
        "A Tabela 2 apresenta recall por mecanismo para os quatro detectores. Quedas, ruído e "
        "picos nominais foram reconhecidos com recall próximo de 1 por todos os métodos, o que "
        "é coerente com a magnitude desses mecanismos em relação ao ruído normal máximo de 0,80 bar. "
        f"A deriva foi mais difícil para a regra robusta ({md(PT['Regra robusta']['deriva'], 3)}) "
        f"do que para LOF, Isolation Forest e One-Class SVM (cerca de 0,93). O sensor travado "
        f"separou os algoritmos: LOF {md(PT['Local Outlier Factor']['travado'], 3)}, regra robusta "
        f"{md(PT['Regra robusta']['travado'], 3)}, One-Class SVM {md(PT['One-Class SVM']['travado'], 3)} "
        f"e Isolation Forest apenas {md(PT['Isolation Forest']['travado'], 3)}. Logo, o padrão de "
        f"baixa variabilidade não é um fracasso exclusivo do Isolation Forest frente ao LOF: a "
        f"regra robusta, que opera sobre o máximo dos escores univariados, captura a planicidade "
        f"quase tão bem quanto o LOF. Dois vetores podem ser normais em amplitude absoluta e "
        f"anormais em estrutura temporal; essa é sobretudo uma demonstração de representação, "
        f"não de superioridade intrínseca de um algoritmo.",
        first=True)
    tabela(doc,
           ["Tipo"] + MET,
           [[lab] + [md(PT[m][t], 3) for m in MET] for t, lab in zip(TIPOS, TIPOS_L)])
    caption(doc, "Tabela 2: Recall por mecanismo de anomalia (média ± desvio-padrão, 20 sementes). Fonte: Elaborado pelo autor (2026).")
    picture(doc, RES / "figura4_mecanismos.png", 16.0)
    caption(doc, "Figura 5: Recall por mecanismo para os quatro detectores. Fonte: Elaborado pelo autor (2026).")

    add(doc, "4.3 Ablação e redundância de atributos", size=12, bold=True, align="left", before=8, after=4)
    add(doc,
        f"A Figura 3 mostra correlação estrutural entre amplitude, máximo de variação, "
        f"desvio-padrão e mudança líquida. A ablação confirma a dependência do sensor travado "
        f"em relação à planicidade: com todos os atributos, o recall de travamento do LOF foi "
        f"{br(ABL['todos']['Local Outlier Factor']['recall_travado']['media'], 3)}; sem "
        f"planicidade, caiu para {br(ABL['sem_planicidade']['Local Outlier Factor']['recall_travado']['media'], 3)} "
        f"(regra robusta: de {br(ABL['todos']['Regra robusta']['recall_travado']['media'], 3)} "
        f"para {br(ABL['sem_planicidade']['Regra robusta']['recall_travado']['media'], 3)}). "
        f"Remover atributos de extremos quase não alterou o F1. Subconjuntos apenas de tendência/"
        f"variabilidade ou apenas de atributos simples reduziram o F1 do LOF para cerca de "
        f"{br(ABL['tendencia_variabilidade']['Local Outlier Factor']['f1']['media'], 3)} e "
        f"{br(ABL['apenas_simples']['Local Outlier Factor']['f1']['media'], 3)}, respectivamente. "
        f"O desempenho no travamento, portanto, decorre em grande parte de um atributo "
        f"explicitamente construído para baixa dinâmica, o que é boa engenharia de atributos, "
        f"mas não deve ser lido como superioridade algorítmica descontextualizada.",
        first=True)

    add(doc, "4.4 Prevalência, PPV e o risco de superestimar a precisão", size=12, bold=True,
        align="left", before=8, after=4)
    add(doc,
        f"O teste possui prevalência de anomalias de 50%, distante do monitoramento contínuo. "
        f"Este ponto é central, e não uma ressalva periférica. Com o recall e o FPR médios do LOF, "
        f"a PPV esperada seria {br(PPV['0.5']['Local Outlier Factor'], 3)} no teste balanceado, "
        f"{br(PPV['0.05']['Local Outlier Factor'], 3)} se 5% das janelas fossem anômalas, "
        f"{br(PPV['0.01']['Local Outlier Factor'], 3)} a 1% e apenas {br(PPV['0.001']['Local Outlier Factor'], 3)} "
        f"a 0,1% (Tabela 3 e Figura 6). Um detector com F1 próximo de 0,97 em teste balanceado "
        f"pode, portanto, produzir maioria de falsos alarmes em produção. Para aplicação real, "
        f"FPR, especificidade e alarmes por 1.000 janelas importam tanto quanto o F1.",
        first=True)
    tabela(doc,
           ["Prevalência", "Regra robusta", "Isolation Forest", "LOF", "One-Class SVM"],
           [[lab,
             br(PPV[k]["Regra robusta"], 3),
             br(PPV[k]["Isolation Forest"], 3),
             br(PPV[k]["Local Outlier Factor"], 3),
             br(PPV[k]["One-Class SVM"], 3)]
            for k, lab in [("0.5", "50% (teste)"), ("0.05", "5%"), ("0.01", "1%"), ("0.001", "0,1%")]])
    caption(doc, "Tabela 3: Precisão positiva esperada (PPV) em diferentes prevalências, usando recall e FPR médios. Fonte: Elaborado pelo autor (2026).")
    picture(doc, RES / "figura6_ppv.png", 15.5)
    caption(doc, "Figura 6: PPV esperada em função da prevalência operacional (eixo em escala logarítmica). Fonte: Elaborado pelo autor (2026).")

    add(doc, "4.5 Severidade, OOD e tamanho de janela", size=12, bold=True,
        align="left", before=8, after=4)
    add(doc,
        f"Na intensidade baixa, o recall cai drasticamente: queda de 1–3 bar foi detectada pelo "
        f"LOF em apenas {br(SEV['baixa']['Local Outlier Factor']['queda']['media'], 3)} das janelas, "
        f"e o travamento de 4 leituras em {br(SEV['baixa']['Local Outlier Factor']['travado']['media'], 3)}. "
        f"Na intensidade alta, os recalls nominais aproximam-se de 1. O benchmark nominal, com "
        f"quedas de 7–25 bar frente a ruído de no máximo 0,80 bar, é relativamente fácil; um "
        f"detector operacional precisa ser testado no limiar entre normal e anormal (Figura 7).",
        first=True)
    add(doc,
        f"No cenário normal OOD, o FPR sobe para {md(OOD['Regra robusta'], 3)} na regra robusta, "
        f"{md(OOD['Local Outlier Factor'], 3)} no LOF, {md(OOD['One-Class SVM'], 3)} na One-Class SVM "
        f"e {md(OOD['Isolation Forest'], 3)} no Isolation Forest. Mudanças operacionais benignas "
        f"fora da faixa de treino, portanto, geram muitos falsos alertas. Isso reforça que o "
        f"experimento principal avalia generalização IID, não robustez a drift de sensor, faixa "
        f"ou calibração.",
        first=True)
    add(doc,
        f"Janelas de 12 leituras (~55 min de extensão) elevaram o F1 do LOF para "
        f"{md(JAN['12']['Local Outlier Factor'])}, ao passo que 48 leituras (~3,9 h) o reduziram "
        f"para {md(JAN['48']['Local Outlier Factor'])}. Anomalias localizadas ocupam fração maior "
        f"de uma janela curta e diluem-se em janelas longas. Mantém-se n = 24 como compromisso "
        f"operacional alinhado à agregação de 5 min do GasTrack, com a ressalva de que o tamanho "
        f"da janela altera o ranqueamento absoluto, embora o LOF tenha permanecido à frente nos "
        f"três tamanhos avaliados.",
        first=True)
    picture(doc, RES / "figura7_severidade.png", 16.4)
    caption(doc, "Figura 7: Recall em função da intensidade da anomalia sintética. Fonte: Elaborado pelo autor (2026).")

    add(doc, "4.6 Sensibilidade de hiperparâmetros", size=12, bold=True, align="left", before=8, after=4)
    add(doc,
        f"O F1 do LOF variou pouco com k: {br(SK['10']['media'])} (k = 10), {br(SK['20']['media'])} "
        f"(k = 20), {br(SK['35']['media'])} (k = 35), {br(SK['50']['media'])} (k = 50) e "
        f"{br(SK['75']['media'])} (k = 75). O Isolation Forest igualmente pouco mudou com "
        f"100, 300 ou 500 árvores (F1 {br(SIF['100']['media'])}, {br(SIF['300']['media'])} e "
        f"{br(SIF['500']['media'])}). O ranqueamento dos métodos não depende de uma escolha "
        f"isolada nessas grades. Os valores k = 35, 300 árvores e ν = 0,01 permanecem os "
        f"pré-definidos; a grade não foi usada para selecionar o modelo publicado.",
        first=True)

    add(doc, "4.7 Janelas sobrepostas, persistência e custo", size=12, bold=True,
        align="left", before=8, after=4)
    add(doc,
        f"Em uma série longa com passo de uma leitura, a concordância entre decisões consecutivas "
        f"foi {br(OVL['Local Outlier Factor']['acordo_consecutivo'], 3)} para todos os métodos: "
        f"alertas são altamente correlacionados. Um único evento de queda persistente gerou "
        f"{OVL['Local Outlier Factor']['alertas_por_janela']} janelas positivas; persistência de "
        f"três janelas consecutivas reduziu para {OVL['Local Outlier Factor']['alertas_com_persistencia_3']}, "
        f"e a combinação com cooldown de duas janelas após o alerta persistente resultou em "
        f"{OVL['Local Outlier Factor']['alertas_com_persistencia_e_cooldown']} disparos, com o "
        f"evento detectado e atraso nulo neste exemplo. Métrica por janela, portanto, não equivale "
        f"a métrica por evento. Propõe-se, para implantação futura — e não como arquitetura já "
        f"pronta —, persistência de 3 janelas, cooldown de 10 min (já usado no GasTrack para "
        f"pressão crítica), reset após 3 janelas consecutivas normais, e escalonamento imediato "
        f"se a regra crítica também disparar.",
        first=True)
    add(doc,
        "A matriz de decisão da Tabela 4 torna concreta a arquitetura de dois níveis. O Nível 1 "
        "permanece baseado em regras determinísticas de segurança (pressão crítica). O Nível 2 "
        "é o detector estatístico de comportamento incomum. Os dois canais são complementares: "
        "modelos de ML não substituem limites validados.",
        first=True)
    tabela(doc,
           ["Regra crítica", "Anomalia ML", "Interpretação"],
           [["não", "não", "Normal"],
            ["sim", "não", "Limite operacional excedido"],
            ["não", "sim", "Investigar comportamento atípico"],
            ["sim", "sim", "Alerta prioritário"]])
    caption(doc, "Tabela 4: Matriz de decisão dos tipos de alerta no GasTrack. Fonte: Elaborado pelo autor (2026).")
    add(doc,
        f"O custo foi medido em CPU hospedeira, não em ESP32. A extração dos onze atributos "
        f"consumiu {br(CUS['tempo_atributos_por_janela_ms'], 2)} ms por janela; a inferência dos "
        f"quatro detectores, {br(CUS['tempo_inferencia_4_metodos_por_janela_ms'], 3)} ms por janela. "
        f"O tamanho serializado foi de {br(CUS['tamanho_kb']['Regra robusta'], 2)} kB para a regra "
        f"robusta, {br(CUS['tamanho_kb']['One-Class SVM'], 1)} kB para a One-Class SVM, "
        f"{br(CUS['tamanho_kb']['Local Outlier Factor']/1024, 2)} MB para o LOF e "
        f"{br(CUS['tamanho_kb']['Isolation Forest']/1024, 2)} MB para o Isolation Forest. LOF e "
        f"Isolation Forest, neste protocolo, não estão prontos para o ESP32: o LOF novelty armazena "
        f"o conjunto de treino. A regra robusta é a candidata natural à borda, pela memória e pela "
        f"explicabilidade. Qualquer afirmação de viabilidade TinyML exigiria medição específica de "
        f"tempo e memória no microcontrolador, no espírito de El Barkani et al. (2024), o que está "
        f"fora do escopo desta prova de conceito.",
        first=True)

    add(doc, "4.8 Limitações e ameaças à validade", size=12, bold=True, align="left", before=8, after=4)
    add(doc,
        "Primeiro, o gerador é altamente controlado. Mesmo após incluir anomalias limítrofes e "
        "um cenário OOD, as classes nominais permanecem mais separáveis do que muitos eventos reais. "
        "Os valores de F1 caracterizam exclusivamente este benchmark. Segundo, o teste balanceado "
        "infla a precisão; a Tabela 3 deve ser a referência para discussão de alarme falso. Terceiro, "
        "não há rótulos causais: o método detecta anormalidade do sinal e não distingue vazamento, "
        "consumo excepcional, falha de válvula ou defeito do sensor. Quarto, treino e teste IID "
        "não avaliam mudança de sensor, sazonalidade, deriva operacional lenta nem calibração "
        "diferente; o OOD aqui é apenas um primeiro indicador de FPR elevado nessas condições. "
        "Quinto, métricas por janela independente subestimam a correlação temporal dos alertas em "
        "janelas deslizantes. Sexto, o custo computacional não foi medido no ESP32.",
        first=True)

    # ---------- 5
    add(doc, "5. CONCLUSÃO", size=12, bold=True, align="left", before=12, after=6)
    add(doc,
        "No benchmark sintético controlado, detectores não supervisionados baseados em atributos "
        "temporais foram capazes de distinguir janelas normais de diferentes mecanismos de "
        "alteração do sinal de pressão. O Local Outlier Factor apresentou desempenho elevado, "
        "sobretudo na identificação de padrões caracterizados por baixa dinâmica temporal, como "
        "o sensor travado, embora sua vantagem deva ser interpretada em conjunto com a "
        "representação de atributos e com a variabilidade entre diferentes realizações do "
        f"experimento (F1 {md(R['Local Outlier Factor']['f1'])}; diferença em relação à regra "
        f"robusta {ic(DIF)}). A regra robusta também apresentou desempenho competitivo e maior "
        "interpretabilidade, além de pegada de memória várias ordens de grandeza menor, constituindo "
        "alternativa relevante para aplicações com recursos computacionais limitados. Os resultados "
        "caracterizam uma prova de conceito de triagem de anomalias e não demonstram capacidade "
        "de diagnosticar vazamentos. A validação em séries reais, naturalmente desbalanceadas e "
        "submetidas a mudanças operacionais, permanece necessária antes da utilização do detector "
        "como componente decisório do GasTrack.",
        first=True)
    add(doc,
        "Como próximos passos, recomenda-se registrar eventos operacionais reais, construir uma "
        "taxonomia de causas, avaliar PPV em fluxo desbalanceado, incorporar métricas por evento "
        "(recall de episódio e atraso de detecção) e medir inferência no ESP32 apenas para "
        "detectores cuja memória seja compatível com o dispositivo. A combinação futura com "
        "sensores de concentração, vazão ou temperatura pode aumentar a capacidade diagnóstica, "
        "sempre mantendo as regras determinísticas de segurança como Nível 1.",
        first=True)

    add(doc, "REFERÊNCIAS BIBLIOGRÁFICAS", size=12, bold=True, align="left", before=14, after=8)
    refs = [
        "BREUNIG, Markus M.; KRIEGEL, Hans-Peter; NG, Raymond T.; SANDER, Jörg. LOF: Identifying Density-Based Local Outliers. In: ACM SIGMOD INTERNATIONAL CONFERENCE ON MANAGEMENT OF DATA, 2000. Proceedings [...]. New York: ACM, 2000. p. 93-104. DOI: https://doi.org/10.1145/335191.335388.",
        "EL BARKANI, Majda; BENAMAR, Nabil; TALEI, Hanae; BAGAA, Miloud. Gas Leakage Detection Using Tiny Machine Learning. Electronics, v. 13, n. 23, art. 4768, 2024. DOI: https://doi.org/10.3390/electronics13234768.",
        "LIU, Fei Tony; TING, Kai Ming; ZHOU, Zhi-Hua. Isolation Forest. In: IEEE INTERNATIONAL CONFERENCE ON DATA MINING, 8., 2008. Proceedings [...]. Pisa: IEEE, 2008. p. 413-422. DOI: https://doi.org/10.1109/ICDM.2008.17.",
        "SCHÖLKOPF, Bernhard et al. Estimating the Support of a High-Dimensional Distribution. Neural Computation, v. 13, n. 7, p. 1443-1471, 2001. DOI: https://doi.org/10.1162/089976601750264965.",
        "SOUZA, Paulo Silas Severo de et al. Detecting abnormal sensors via machine learning: An IoT farming WSN-based architecture case study. Measurement, v. 164, art. 108042, 2020. DOI: https://doi.org/10.1016/j.measurement.2020.108042.",
        "TAN, Qiong et al. A new sensor fault diagnosis method for gas leakage monitoring based on the naive Bayes and probabilistic neural network classifier. Measurement, v. 194, art. 111037, 2022. DOI: https://doi.org/10.1016/j.measurement.2022.111037.",
        "TSOUKAS, Vasileios et al. A Gas Leakage Detection Device Based on the Technology of TinyML. Technologies, v. 11, n. 2, art. 45, 2023. DOI: https://doi.org/10.3390/technologies11020045.",
        "YUAN, Hongyong et al. Real-time detection of urban gas pipeline leakage based on machine learning of IoT time-series data. Measurement, v. 242, part B, art. 115937, 2025. DOI: https://doi.org/10.1016/j.measurement.2024.115937.",
    ]
    for ref in refs:
        add(doc, ref, size=11, align="justify", before=0, after=6, first=False, line=1.15)

    doc.save(OUT_ART)
    print("artigo", OUT_ART)


def construir_carta():
    doc = Document()
    setup(doc)
    add(doc, "CARTA DE RESPOSTA AOS REVISORES", size=14, bold=True, align="center", after=4, line=1.15)
    add(doc, "Artigo 2 — Benchmark sintético de prova de conceito para detecção não supervisionada de anomalias (GasTrack)",
        size=12, italic=True, align="center", after=8, line=1.15)
    add(doc,
        "Prezados editores e avaliadores, agradecemos a leitura cuidadosa. Concordamos com a "
        "decisão de reexecução do benchmark. A versão revisada remove a amplitude duplicada, "
        "separa treino/calibração/teste, repete o experimento em 20 sementes e reescreve "
        "Resumo, Abstract, Resumen, Resultados e Conclusão. Indicamos abaixo, ponto a ponto, "
        "as alterações. O código está em experimento_anomalias.py e os números em "
        "resultados_anomalias/metricas.json.",
        first=False)

    itens = [
        ("1. Amplitude duplicada",
         "Removida. O vetor passou de 12 para 11 atributos. Todos os modelos foram reexecutados."),
        ("2. Ablação",
         f"Concluída (5 sementes). Sem planicidade, o recall de sensor travado do LOF cai de "
         f"{br(ABL['todos']['Local Outlier Factor']['recall_travado']['media'], 3)} para "
         f"{br(ABL['sem_planicidade']['Local Outlier Factor']['recall_travado']['media'], 3)} "
         f"(Seção 4.3)."),
        ("3–4. Planicidade e sensor travado",
         "Discutido como estudo de caso de representação. A regra robusta empata praticamente "
         f"com o LOF no travamento ({md(PT['Regra robusta']['travado'], 3)} vs "
         f"{md(PT['Local Outlier Factor']['travado'], 3)}); o Isolation Forest permanece próximo de zero."),
        ("5. Calibração independente do LOF",
         "Protocolo 5.000 treino + 2.000 calibração + teste. Com novelty=True, score_samples "
         "nunca é aplicado ao conjunto de ajuste (Seção 3.5)."),
        ("6. contamination vs limiar externo",
         "contamination='auto' e decisão exclusivamente pelo percentil 99 dos escores na calibração. "
         "Orientação unificada: maior escore = mais anômalo."),
        ("7. One-Class SVM",
         "kernel='rbf', nu=0.01, gamma='scale' (γ = 1/(n_features·Var(X))). Python "
         f"{M['hiperparametros']['python']}, scikit-learn {M['hiperparametros']['sklearn']}."),
        ("8. random_state e sementes",
         "Isolation Forest usa random_state=semente. Gerador numpy.random.default_rng(semente), "
         "sementes 0–19. A subseção de reprodutibilidade agora descreve o experimento, não apenas o raciocínio."),
        ("9–10. Múltiplas sementes e comparação pareada",
         f"20 sementes com média, dp, mediana e IC95%. Diferença de F1 LOF−regra = {ic(DIF)}. "
         f"McNemar na semente 0: n01={MCN['n01']}, n10={MCN['n10']}, p≈{MCN['p']:.1e}."),
        ("11. Coerência aritmética",
         "Mantida: o recall global é a média dos cinco mecanismos (500 casos cada)."),
        ("12–13. Teste balanceado, FPR e PPV",
         f"Mensagem central da Seção 4.4 e Figura 6. FPR do LOF {md(R['Local Outlier Factor']['fpr'])}; "
         f"PPV a 1% = {br(PPV['0.01']['Local Outlier Factor'], 3)}; a 0,1% = {br(PPV['0.001']['Local Outlier Factor'], 3)}. "
         "Especificidade e alarmes/1.000 na Tabela 1 e no texto."),
        ("14. PR-AUC",
         f"Incluída na Tabela 1 (LOF {md(R['Local Outlier Factor']['pr_auc'])})."),
        ("15–17. Gerador, código e reprodutibilidade",
         "Equações e distribuições na Seção 3.2–3.3. Código, seeds e requirements-anomalias.txt "
         "anexados. A formulação “reprodutibilidade do raciocínio” foi substituída."),
        ("18. OOD",
         f"Cenário normal fora da distribuição: FPR LOF {md(OOD['Local Outlier Factor'], 3)} "
         f"(Seção 4.5)."),
        ("19–20. Anomalias limítrofes e severidade",
         f"Três intensidades. Na baixa, recall LOF de queda = "
         f"{br(SEV['baixa']['Local Outlier Factor']['queda']['media'], 3)} e travado = "
         f"{br(SEV['baixa']['Local Outlier Factor']['travado']['media'], 3)} (Figura 7)."),
        ("21–23. Atributos, correlação e padronização",
         "Onze atributos com fórmulas (Seção 3.4). Matriz de correlação na Figura 3. "
         "StandardScaler só no treino, aplicado a calibração/teste sem reajuste."),
        ("24–27. k, ν e sensibilidade; sem vazamento do teste",
         "k=35 justificado a priori (ponto médio 10–50 de Breunig). ν=0,01 = orçamento de 1% de FPR. "
         f"Sensibilidade: k=10…75 mantém F1 entre {br(SK['10']['media'])} e {br(SK['75']['media'])}. "
         "Hiperparâmetros não foram escolhidos no teste."),
        ("28–29. Janela de 24 leituras",
         "Clarificado: 23×5 = 115 min (aproximadamente duas horas). Justificada pelo dashboard "
         f"de 5 min e pelo compromisso latência/contexto. F1 LOF em n=12, 24, 48: "
         f"{br(JAN['12']['Local Outlier Factor']['media'])}, "
         f"{br(JAN['24']['Local Outlier Factor']['media'])}, "
         f"{br(JAN['48']['Local Outlier Factor']['media'])}."),
        ("30–31. Sobreposição e métricas por evento",
         f"Série deslizante: acordo consecutivo {br(OVL['Local Outlier Factor']['acordo_consecutivo'], 3)}; "
         f"{OVL['Local Outlier Factor']['alertas_por_janela']} janelas vs "
         f"{OVL['Local Outlier Factor']['alertas_com_persistencia_e_cooldown']} alertas após "
         "persistência+cooldown. Distinção janela/evento explicitada."),
        ("32–33. Definição de travado e deriva",
         "Travado: k=8, σ=0,02 bar. Deriva: pᵢ ← pᵢ − D (tᵢ/t_fim)², D~U(8,20)."),
        ("34–35. Quatro detectores e regra robusta",
         "Tabela 2 completa. A regra robusta recebe destaque operacional (F1 próximo, memória "
         f"{br(CUS['tamanho_kb']['Regra robusta'], 2)} kB, explicabilidade)."),
        ("36–37. Custo e TinyML/pressão",
         f"Tempos e tamanhos na Seção 4.7. LOF ≈ {br(CUS['tamanho_kb']['Local Outlier Factor']/1024, 2)} MB: "
         "não recomendado para ESP32 sem nova avaliação. El Barkani et al. citados como evidência "
         "de inferência embarcada, não de transferibilidade (câmera térmica, não pressão)."),
        ("38–39. Caracterização do GasTrack e faixas sintéticas",
         "XDB305 0–250 bar, 4–20 mA, ESP32-ETH01, ADC 12 bits, MQTT 10 s, agregação 5/15 min, "
         "crítico 20%, cooldown 600 s. Faixas 60–145 bar e 0,5–5 bar/h declaradas sintéticas."),
        ("40–42. Anomalia ≠ vazamento; dois níveis",
         "Mantido e reforçado. Título e conclusão usam “anomaly detector for pressure telemetry”. "
         "ML não substitui limites de segurança (Tabela 4)."),
        ("43–44. Persistência/cooldown e matriz de decisão",
         "Parametrizado: persistência=3, cooldown=10 min, reset=3 janelas normais, escalonamento "
         "se a regra crítica disparar. Tabela 4 inserida."),
        ("45. Marcadores de template",
         "Não estavam presentes nesta cópia do manuscrito; a versão revisada também não os contém."),
        ("46–47. Referências",
         "Removidas Gubbi et al. (2013) e Viegas Junior (2026), não utilizadas no argumento central. "
         "Todas as oito referências restantes são citadas no corpo."),
        ("Avaliador 2 / título / conclusão",
         "O trabalho é declarado prova de conceito sintética no título, no resumo, na metodologia "
         "e na conclusão. A formulação sugerida para a conclusão foi adotada, com os novos números."),
    ]
    for titulo, corpo in itens:
        add(doc, titulo, size=12, bold=True, align="left", before=8, after=2, line=1.15)
        add(doc, corpo, size=12, align="justify", first=False, after=4, line=1.5)

    add(doc, "Materiais anexos", size=12, bold=True, align="left", before=10, after=4)
    add(doc,
        "experimento_anomalias.py (gerador, modelos, sementes); requirements-anomalias.txt; "
        "resultados_anomalias/metricas.json (todas as repetições agregadas); figuras 2–7; "
        "matrizes de confusão da semente 0 no JSON (campo confusao_semente0).",
        first=False)
    add(doc, "Atenciosamente,", size=12, align="left", before=12, after=2, line=1.15)
    add(doc, "Os autores", size=12, align="left", after=0, line=1.15)
    doc.save(OUT_CARTA)
    print("carta", OUT_CARTA)


if __name__ == "__main__":
    construir_artigo()
    construir_carta()
    for dst in (DESKTOP,):
        if dst.exists():
            shutil.copy2(OUT_ART, dst / OUT_ART.name)
            shutil.copy2(OUT_CARTA, dst / OUT_CARTA.name)
            print("copiado para", dst)
